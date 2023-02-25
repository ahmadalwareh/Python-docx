from docx.api import Document
doc = Document()

r = doc.add_paragraph().add_run()

if r.instr_text is None:
    print("is none")
    r.instr_text = "fee"
    for el in list(r._r):
        print(el.tag)
